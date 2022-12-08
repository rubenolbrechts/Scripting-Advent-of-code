#!/usr/bin/python3
################################################################################
# Word files (example-code-5-docx.py)
################################################################################
# sudo dnf install python3-docx
from docx import Document
# Opening a blank document based on default template
document = Document()
# Body text is divided into sections
# each of which starts with a heading
document.add_heading('My first Word document in Python', 0)
# Paragraphs are used for body text
# but also for headings and list items like bullets
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
################################################################################
# Adding more headings and text
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
################################################################################
# Adding an image
# default appears at native size calculated as pixels/dpi
from docx.shared import Inches, Cm
# document.add_picture('python-scripting.png', width=Inches(1.25))
document.add_picture('python-scripting.png', width=Cm(5))
################################################################################
# Adding a page break
document.add_page_break()
# Heading level 2
document.add_heading('Table of genes', level=2)
# Add table
table = document.add_table(rows=1, cols=3)
# Word has pre-formatted table styles
# Hover mouse over thumbnail in Word table style gallery
# Style name is formed by removing all spaces from table style name
table.style = 'LightShading-Accent1'
# Table data
genes = (
    (7124, 'TNF', 'tumor necrosis factor'),
    (4049, 'LTA', 'lymphotoxin alpha'),
    (4050, 'LTB', 'lymphotoxin beta'),
    (7132, 'TNFRSF1A', 'TNF receptor superfamily member 1A'),
    (8743, 'TNFSF10', 'TNF superfamily member 10')
)
# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Gene ID'
hdr_cells[1].text = 'Symbol'
hdr_cells[2].text = 'Description'
# Add data rows
for gene_id, symbol, desc in genes:
    row_cells = table.add_row().cells
    row_cells[0].text = gene_id
    row_cells[1].text = symbol
    row_cells[2].text = desc
# Save document
document.save('example.docx')
################################################################################